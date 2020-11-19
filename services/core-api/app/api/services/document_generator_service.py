import requests, hashlib, os, mimetypes, json, datetime
from flask import Response, current_app, stream_with_context
from app.config import Config
from shutil import copyfile
import os
import docx
import base64, io
import re


def sha256_checksum(filename, block_size=65536):
    sha256 = hashlib.sha256()
    with open(filename, 'rb') as f:
        for block in iter(lambda: f.read(block_size), b''):
            sha256.update(block)
    return sha256.hexdigest()


class DocumentGeneratorService():
    document_generator_url = f'{Config.DOCUMENT_GENERATOR_URL}/template'

    @classmethod
    def generate_document_and_stream_response(cls, template_file_path, data):

        # Ensure that the desired template exists
        current_app.logger.debug(f'CHECKING TEMPLATE at {template_file_path}')
        template_exists = False                  #cls._check_remote_template(template_file_path)
        if not template_exists:
            current_app.logger.debug(f'PUSHING TEMPLATE at {template_file_path}')
            template_file_path = cls._push_template(template_file_path)

        # Create the document generation request
        file_sha = sha256_checksum(template_file_path)
        file_name = os.path.basename(template_file_path)
        file_name_no_ext = '.'.join(file_name.split('.')[:-1])
        # https://carbone.io/api-reference.html#native-api
        body = {
            'data': data,
            'options': {
                'reportName': f'{file_name_no_ext}-{datetime.date.today().strftime("%d%m%Y")}.pdf',
                'convertTo': 'pdf'
            }
        }

        # Send the document generation request and return the response
        resp = requests.post(
            url=f'{cls.document_generator_url}/{file_sha}/render',
            data=json.dumps(body),
            headers={'Content-Type': 'application/json'})
        if resp.status_code != 200:
            current_app.logger.warn(f'Docgen-api/generate replied with {str(resp.content)}')

        return resp

    # TODO: This logic should be moved to the document manager?
    @classmethod
    def _push_template(cls, template_file_path):
        current_app.logger.info(f'****************---------------****************')
        current_app.logger.info(f'template_file_path:{template_file_path}')

        file = None
        file_name = None
        pathz = None
        if template_file_path == '/app/app/templates/permit/Permit Template.docx':

            idx = template_file_path.index('Permit Template')
            new_path = template_file_path[:idx] + 'test_' + template_file_path[idx:]
            current_app.logger.info(f'new_path:{new_path}')

            copyfile(template_file_path, new_path)
            cls._prepare_template(new_path)
            file = open(new_path, 'rb')
            file_name = os.path.basename(new_path)
            pathz = new_path
        else:
            file = open(template_file_path, 'rb')
            file_name = os.path.basename(template_file_path)
            pathz = template_file_path

        files = {'template': (file_name, file.read(), mimetypes.guess_type(file_name))}
        resp = requests.post(url=cls.document_generator_url, files=files)
        if resp.status_code != 200:
            current_app.logger.warn(f'Docgen-api/push-template replied with {str(resp.text)}')
            return pathz
        return pathz

    # TODO: This logic should be moved to the document manager?
    @classmethod
    def _check_remote_template(cls, template_file_path):
        file_sha = sha256_checksum(template_file_path)
        resp = requests.get(url=f'{cls.document_generator_url}/{file_sha}')
        if resp.status_code != 200:
            current_app.logger.warn(f'Docgen-api/check-template replied with {str(resp.content)}')
            return False
        return True

    # TODO: This logic should be moved to the document manager?
    @classmethod
    def _prepare_template(cls, template_file_path):
        current_app.logger.info(f'********************************')

        file = open(template_file_path, 'rb')
        file_name = os.path.basename(template_file_path)

        photob64 = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAfQAAAFuCAYAAAB3K+qvAAAAAXNSR0IArs4c6QAAAARnQU1BAACxjwv8YQUAAAAJcEhZcwAAFiUAABYlAUlSJPAAADdLSURBVHhe7d1PqGVVvtjxbROIwkta4YVXTQLambSO2oLwunukkkG3EHjt6JWjtkdVBYGuGnUVBEpH6iRVjiwzUUflTBsC1Q2BUhJQM1FHVkOICi+UwgP7hUA5q5zvuXtdf3fXPveeP3vvs9da3w9sPfdP3XPOPnvv31q/9VtrP3BvoZEkSVn7Qft/SZKUMQO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6pOJdv369eeaZZ5pXXnml/Y5UngfuLbSPJakon376afPaa681b7311vLrBx98sHn99debF154Yfm1VBIDuqTivPPOO83ly5ebL7/8sv3O9wjqd+/ebb+SymFAl1QUAvlJqXUveyqRY+iSivDdd981zz///JFgfurUqebcuXPtV1LZ7KFLyt5f/vKX5tlnn20++uij9jtN8+tf/7p58803m4cffrh54IEH2u8e8LKnEtlDl5Q1xsl/8YtfHAnmFy5caN59991lMO967LHH2kdSWQzokrJFED99+nRz+/bt9jvNsor96tWr7Vf3O3v2bPtIKosBXVKWqGRnbjnpdtAbv3nz5n1j5kxdi5yyplIZ0CVl59q1a8sCOArhQBr9ww8/bH71q18tv47efvvt9tEBCuWkElkUJykrFy9eXAb05Oc///myZ943Xo5HHnnksBcPL3kqlT10SVlI09JiMKeS/datWyuDOen2GMylkhnQJc1empbGuHnCWDmV7Kz8tsr777/fPpLKZ0CXNGtMS6P4LQbnl19+eVnNfpIPPvigfSSVzzF0SbNFyvy55547XJN905urdMfP4SVPpbKHLmmW6JHTM0/BnHFyUuzrBnPHz1UbA7qk2WGsnDHzFJCZakbxW9+0tFUcP1dtDOiSZqU7x/zxxx9fzjF/8sknl1+vy/Fz1cYxdEmzsekc81W+/vrr5sc//vFhoyDykqdS2UPPVBpffOutt9rvSPnaZo75cc6fP38YzDft2Uu5soeeoW7vw49QOUtzzOPd0phjvs60tD7vvffesjI+IV3P3dgSzxeVyh56hl599VV7HxMiC0IDintqx43vkSKmgaXt9N36dN055n04L+idJzQMSNtLVaCHrnzcvXv33qlTp+hiLLd33323/YnGcOXKlcN9vWp78MEH7y2CUPsvtK5PPvnkyLHMfnzzzTfbn27n0qVLh3+Pv33nzp3l99P32KRSeXRn5saNG4cXpscee6z9rsawTjCPm0F9fTdv3rz38MMPH+47HvO9XdBAoFGQ/mZsHKTvsUmlMuWemTfeeKN91DRnz55tH2lopG1feuml9qtmOf/57t27y/HXtLHISRzyiL+v1RjC2HWOeZ9YCPf0009733NVx6K4jNy+fbt54oknlo9ZAvOLL77w3s4DI8hQUBUXJSHQrLoJCAHkoYcear9adv/aR+rz4osvHmn4MMecaWncz3wXNBJ++9vfLh/zOXXnrVPzkPgZqVT20DMSe+cEGYP5sGgwUaAVgzm9vOPu6LXq+7ofATcGc4rVCLy7BnOKEi9fvtx+1TQXLlywWFRVMqBnhOk4ien2YRHECeYE9YRq6zfffNOgvaM0LS2umbDLHPMugnmaaUDj4MqVK8vHUm1MuWeE1K5zz4dHoInjrwRwAvmZM2eWX5/EdO5qBFqCOTdKSXaZY97FdLc4x5xsCo2FLj8j1cCAnok4Vkuq/c6dO8vH2g0BgRX3UjBn3zKmu0nK1oZWP7IdBPN0tzSQ9bh06VL71W7Y5wTz1FggkBPQ+xjQVQNT7pmIi5cMkabUQUBgXDcF421vAmItw/1Szznex5ysx1DBHCwTm4I5f3+oXr/yc/369WXD/JVXXmm/UycDeiZiQDeADIM0exozJyDQu9u1QEsHtz7l4pqmpdEAZd8OOY2MhkIssKPn73lRJ2ZOcC5TB1P71FEDeibSxRFeuHbHuHks0qJ3Rw9du2GfxlufcqwOMce8K9Y8kFFhXF71oWceg3g6JmplQM+EKffh0LtjDfaEsVcXIdkdwTzNBce2Qxgn4Xn++Mc/tl81zkSoFNfEeB7LgJ6NGND/5m/+pn2kTdGCpweZMh6k2AkIu4hFX7XqBnOC+BBzzLs4D7aZc27AL0+8SZUOGNAz8U//9E/tI1PuuyA9F+/sdePGjZ0yHrGhVev4e18wH2qOeVd3zjlj5yfh99OFP+fsFo1QxoupT2Cf14wMTbx3vg4Y0DMRA4cp9+1QNBOrYFmAZNdba8beeY0BfcpgzkU8BjLqHtbpeccG3NDp/ylQyU9qmdv10iDlOI5ZitqwP8iyJd4e93sG9EzEgG4PfXPsv3gR4OYd9HZ2VfPnMmUwp4dNIVzCoj/rFtrFRW1yCujsX4L46dOnl73RWBjLcRffVy1439xrIQ6ZkWXTAQN6Jqxy3w2BJwVfAs5QF4Fae+hTBnPQM037muP/6tWry8fr+Oyzz9pHTfOzn/2sfTRvNF7Yv8fVZ9Q2RYtGXVyoyKmm9zOgZyKe2KbcN0PvplsVPVSj6Jtvvmkf1VOsOHUwpycah0o2nXOeU8qdhjtj5EzHSnivTMtjH3/yySftdw/u7VBTL50MW3y/BPOcMi6TYOlXzdu3337LWpXLbXHRbL+rdXz++ef3Fi35w/134cKF9ifDOHPmzOHfXlxg2u+Wa9EYOny/bIsL6vL4HBPPkZ7v6aefbr+7njt37hz+27mfO4tgfW/R2zx8vWy//vWv7929e7f9jQN8L/68Bpy3cb9cvXq1/cmB+LOaGdAz8OGHHx4erD//+c/b72od7K+07x5//PH7Lo67IsCkv8/nVLJ9BPNFb/zw+WiYffHFF+1P1kMjK/37TRsDU+J10uBIr5XtypUr7U+PIvDH3yvd66+/fuT99jXK489rZkDPQDygz507135XJ+kGAy6EQ6L3F3v/fF2qfQRzGkhx//J5boqgmP790NmZocTXyMZ7PinbE3+/ZDdv3jxyDPzqV79qf3JULfvjJAb0DMR0UzfVpH7dVPuq3s4uaFylv0+AK9U+gjk98VOnTh0+J5mWbbIrMT1948aN9rvzwPuJQzZspNzXaXjGf1MqzuGYteC4W3UM1LA/1mFAzwCt0nSw0mLVyWKq/bgLwba6ac9bt261PykLDcj4PqcI5nxWPE96TgL7ttmP2CggQMwFDZb4HtkYElh338Z/VyL2Q6wn4HM8bril9P2xLgN6BmIrteS07lDGTrWjhsKkbip4imCOuG/5/LatTZhrQRzvpzteTrZnk0Zn/LelYT/E2pR1zuGS98cmDOgZ8GBd3xSpdrIk8TMZo8Gwb3E4gY2MxxTB/NKlS0eed5c0+RwL4hi+iMcnj6mR2VT692yl6Q5DrDN7JP5+zYwQGfBgXd/YqXbEVOlcC6128cILLxy+PzaGfMbYj10xALMR3Hcxt4K4bkOQXvq2QzXx75QkZtfY1i2EjP+mZkaIDHiwrmeKVHssEOM5ShoCIWjHeg02ektTBHM+q9hzXVXNvIk5FcR1K/ZpFG46BS9Kf4etFHxG8X3RsFxX/Hc1M0JkwIP1ZPR04gVzjFQ7gS0WWW0zjWquSKd3g/kmF9Rd0CiKBVCsFzBEen8uBXHdam3e664NwfS32ErQbfAwRLJJQ7K0/bEtI0QGPFiPx8UxXry3neJ0kljxzfNN0XOdAsGzW3G9a7p7XezDOExC4Nul55rMpSCO9xIbK0O9v/T32HLH/ojnL/tr0wZdSftjF0aIDHiwHi9WxHJhGOKC2UWAiD2IbQqZ5oh91Q3mY2Q3VukWQA01/W8OBXHdhhLHz1DDQOlvsuWsu49o8GyTTSllf+zKCJEBD9bV4qI7bGPNB4/PwwWoBN3eIxs1AlPpznEfctGkfRfEkXnoTr0acg2J9HfZchaHeXbZR6Xsj10ZITLgwdqvWxU91ph2t3e+zjSauaOnGNOcvL8pg3m34nvoJY33XRAXn59t6H0bj8dcdRvju2S94t+pmREiAx6s96N3GQuNxlzcpbTeOcE87ruhe48n6RaJbVoAtY7YWJm6IK47h3/ohiZp6vS32Y85WueGK5uIf6tmRogMeLAexcWfSui0T7YpollXab1zAnd8PwSEKZet5XOKaX4es4+HtM+CuJjqZxsj3U9FePr7FBTmplvRPsQUxfS32GpmhMiAB+tRceETLgxj9sBK6p3TGOkG8zHm6q/SN648xvPHdP6UAS+uUcBGwd8Ycr77YrcxPtQtjdPfY6vZDxY7QMrG9evXm7feeqv9qmkWF7dmcVFovxrW119/vXy+ZNH7ah/lh3323HPPNd99993y60XPuFn0lJpFI2X59RQuX77cvP/+++1XTXPjxo1Rnv/TTz9tHzXNIqC3j8Z1+/bt5vz58+1XTbPodTaLAN9+Naw///nP7aOm+clPftI+ygP7iH2FRYOuWTQyl//XQNrArhnjY0pbzbqpujHSmVEpvfNuNTlp7jGm9h2nO2Y65tS4WJA2VaFfnEvPsTLWEBBiZfiUtQ+76q4EN+RnE/9uzQzoGfBgvX/slQvo0IVUUSlj590x3bGDTR+GROK+HLOAETGlO8WQQnfJ4bGL8GLB39D1B2MZu4g1/V22mhnQM+DBOs3iMVEJvfNuMKcRNHUwp9HF/kuvYagx01V4f+m5CK5jPhe6jZWxlwOO7y+XCnc+g5jBoGE+9HGY/jZbzQzoGaj9YO0GprGrskvonXf3GWnasYNbn/g6pui9cmyk55uiIRYDFY/HlmOFe/cY4D0MLf19tpoZ0DNQ88HaXYBkihui5N47n0swj8GHbciV4FaJ9QJjV4BPnWpHbhXusYHFxrE5hvgcNTOgZ6DWg3XKxWOSbuFdbr3zuQRznjPWPEy1nnqc0jhmA2LqVHsSG5tTNJB2QVo9jvePeQyk52CrmQE9AzUerFOMu3WRau8W3uVkLsEcMbDSKJuqeCuO14+R2k2mTrUnOVW4x9kGBPYxj4H0PGw1M6BnoMaDNS6fOVU6c+rCuyHNKZiT1YivZaosB+83Pu9Y738fqfYklwr37lTJsY+B+Fw1M6BnIKb2xu6lzkF3xa0p5hJ3A+KUy6Huak7BnCATgw499akwRS09L9X0Y9hXqh25VLhzDMR9NPZ6EUjPxVYzA3oGYs9xqoUy9oWLQRw3n+JiMFXhzhjmFMwRU8IMX0z5WmLB2FjLrsZzceohmVwq3GOqnSGQKY6B9HxsNTOgZyCmr8a6UM0F7y+9Vy5aY18MuoV3UxVvDWFuwTymotnGHMPuE4dpxigYiw2GqVPtyKHCvTvcMtUxEJ+zZgb0DBB00sFK8NnnRXtMcYoaF8yxV/liP8bipikK74Yyt2Aeq6/Z9pHliJ/l0AVjnIP7SrUnc69w5/iLwy1TNjrSc7LVzICeibicZU7ju+viYhArzC9dutT+ZDzxAsnFeuoe5bbmFMxpAMU0O9sUmZU+MeAO3TDbZ6o9ift5jhXunLPp9Y1d1d6VnpetZgb0TMSTZYpx5al1LwZjB4RuanDuc3qT7k1O9hnM6bXGhiYb46f7eD2xII6G4ZD2nWpP5lzhzv6PDaqpa33S87LVzICeiVi4NfQFa9/ixZht7Cku3UrlKRasGUL3de8zmHM8xtoDtimyKqvEO3kN+XnOIdUOPuf0Gng9cxOHO/ZRh5Kem61mBvSMxAvovnoJY4iLgYwdXLkwxl7l2DcLGQqvMV4095XWBj3WGOR4TEDdpzh8MuT4fXef70ts9M5tOeI4zZRjYezalz7p+dlq9oPFDlAmFsGufdQ07733Xvsob9euXWs+/fTT5ePFxaBZBIvl47GcP3++uX379vIxz/fuu+8u/z93r7zySvPRRx8tH/N6FxfRvbzuixcvLvfhd999t/z61KlTzaK33pw5c2b59b6kYwiLgNc+2k3fPt+XL7/8sn3UNI899lj7aP++/vrr5vLly+1XTbNoWA22/7WFNrArAzGtmNP0qlW6C1CMPY7dHX/OZU5/nH/MxvuYWl/xGz1FUtJzELNXQ4wvd8eE911jEacE7nNooysu8ctQ4L6yRuk1sNXMgJ4RLqrxIjO3wphNkV5P72XsNGK3CC6HO1WBC2Ss/ieoTm1OxW99eH3pdRHYd8X74nhMf3MOjec4x34fDbo+3Ybm2LUvx4mvo2YG9MzEXtJcTuxtdAPsmONuTPGJDaF9jj9vKvaApp4KhLkVv/WJx9IQDZ44LZDjZg5ZiDhtbg7TVruNnrFrX06SXgdbzQzomYkFKLmm3bkYxCk4Y07D66ZO6e3msnhMt9EzdQ+IBmPcdzzed/FbnzjlcdfGRrfXOZdhmblNWYtDABwX+35N6bWw1cyAnpkS0u6xIpkL1Vi9ZYJ57F0SzOcy5nsSPtd4EZ/yJieInxEbr2WuC+/EoZtdGhwch3FoYR/DG314Xek1ce7vG8dmvAbNYQ2H9FrYamZAz1DOaXeCbDz5xup1Erjj2DOBPaepfvEznrLYaO7Fb31iw2eXzziOU/M359JYjufM2LUm6+jWvkx1bB4nvR62mhnQM5Rz2n2KcTcuxN1gPuYY/dBopKXXzjZVz5igPefitz685vRadymIiws3sc1paCEOvex7rLo7DDSXrE18TTUzoGco17R7d9xtjF4f+yY2GnieOa57vUp3NbipbnKSQ/FbnziVc9sUOcdMbADO7Y6Gc5myRsMuZkPmNFMkvSa2mhnQM5Vb2r0bqMZYQpMLTlzZiy2nYN59/VNV45PxiZ8Nj+dY/NZniBXiujMJCPBzMpcpa7H4cE5DEkivi61mBvRM5ZR2JyjFXvMY4278ve7Y71wqlNfVnS41xZh/N4XKhXquxW99YgNom8Zb9/3PsQE4hylr3dkiczu30utiq5kBPVPdtPucx4hjy36sQEWaND0HW27Fgt3pUlO8/u78fBpacy5+66IRF1//pj1rephzTSFH8TXuq1ccG05z7ECk18ZWMwN6xmKqkBNujrrFRmMEqrgf2KYadx5Kdwx3isInGhAxGFIMN7dU80liI4jXv6nYCJxyJsEmeE3pNfJ57UPMBvIa5th5SK+PrWYG9IzRWo8X5bGmgG2rG6i2LVo6TkxTs425SM1Ypl4NjgtyrvPzI+Y/p/ew6Tz9GKTY5rD6Wh8+q/QayaBMrZvFmGuhZHp9bDUzoGcuVsBy4s2plxF7QGMEqvje2Ta9qM9BrNJmG3sMl8AdgzmfS47BHPH42mRxE47DuA/mXM2/7ylrsbE51ywG0mtkq5kBPXOcYLEFPUb1+Da6PaChAxWp+/j3x+j9j60bWMYewyVwx4wJzz3n2ouTxPeyyfuIRWak6ucapLDPKWvduo65ZQCj+DprZkAvQGzFk4IfO2V7km4vcOhAFafxsBHM53xRXmXKwNId/pjrWOi6OMbje1lXbAjmsA/2NWWNYzHOTNn3gjYnSa+TrWYG9ELEKtR9p57HClQEpfi32bjo5FbMhTj+S2AZc6oY+ydenHm+OU7P2kRsxK5bdb2vRXt2EY/3Kcf5Y2aAfbbvTsJJ0mtlq5kBvRCxeIZtXxer7oVgqEDF+4s9TLa5L0u6ChfmqQIL+yc29thyD+aIUyHXKYTs7gce53DsxOG0qYIqGbZ4fM7h5isnSa+VrWYG9ILEAha2qYN6dyrUUOP59Mbi393HexsKF8t4kR5zuiEBK/fFdlaJPdd1xnY5XtLvcyxNsWjPrmIjncbsVOIxQ2Ynh4ZPer1sNTOgF6TvAj5V4OO5Sa+n5103DXqSeCFmY2x+zsU5x2EfxV7iGJX/ERmMuO+mHIMdW2zgnbQPu2sh5NKo2WVa3ra6sy7GHAoaUnzNNTOgF2ZfQT0W7xB0d50KxfvoBiR6KTkXcsUMCgFpzIvlvrM1Y9qk59otBpx7cVcUsxBTrK3PvorZo7FnXQwpvWa2mhnQC9QX1BlnHKt4jDHZ+Fy7XnxoDMQiLjYubjkWvyWxt8U2Zm853rCEbc7zrLcRK9VPCtCxYTN2RmRInMMxCzHFsR8b5eyrnM639LrZamZAL1RfUOckHbqlzwUytup3TQ2SHo1T3thyXP0t6hbBjdnz6Q5RTJWqnVIM0scVbE29aM+QYhU/jduxkS2K+2qKjMCQ4muvmQG9YH1pazZ6u0MVBcW/T2pzl1Z9rJBnIwjmPu7bVwTH5zKGbhZgbvf1Hkqs1Vg1bMF+jw3D3BqFMcsy9nAJx2PMiNERyE167Ww1M6BXgNZ2DCpp46Kxy1h3tze47TxZGhexWIyN15tLQc4qXCinKoLrrszHRXmshsM+0WBM75EG36r3GMef574aXJ847r/tebWu7lTTXetf9iG9fraaGdArwYUwzt2NGz25TYNnN5hv24vgYhLT0WwEwVzGOo8TU8O8x7EaKN0xc4JZicEc66SiuwEqt0JKAmp6/WQZxvwsea54/g011XRq6fWz1cyAXhl6w7H3EjcC6TpjZ91gvk1vsK9XzoUl1wtK1xRFcDTSunUSBLmcipk2FY+9vjQ6wTv3ABWPnbGr8uPxw7GTa0MwvQe2mhnQK0Uar298nY10HxeVvsAwRDBf1SvPYbGPdUxRBMe+imPJbGRaSu2ZJzEAdRufvPcx1kKYWjwvjyv621W3aDDnIa74PmpmQK8cKTd6Ot3Kcja+RzBKgXbXYF56rxzsz7GL4GgwdD+vbYc8chPfd3esN0674vdyHLbhWImNwe57HAqN9XicjjnzYgrpfbDVzICuJU5wegOxGCdu3d7gScGcCxGBh79JsKG3VHKvHOyP2GAZowgupmPZ2Ke5rpy3qbigDAE7GnothH3hnEnvgXNxLLHxw3Ga+zBNei9sNTOg6z5cDLs96XW3bs+xbyutV56MWQRHYyH+fTYu+DmvnLepWMkfx5aHXgthn2Lh6li9Zo7L9BxsuTZ+ovh+amZA10qc+Kt67NtupfXKk27PecgiOAJWt4HF1yUXv/WJvco4xDDkWgj7FueDj5F5oWGY+5zzPun9sNXMgK6VumPm6/S+08aFlTQ74/P8HVKJpfYmaaDE4YQhe1bss26jir9/3HBHqWIg4nhCXAaWLefCLoap4nsZo2FCZiz9fY7Zscbopxb3W80M6OrVDebbVLPXIvaeh5z6Qyo0NhTYxqx6njOCW9wPfM1+jqn22GvPURxSGKNCn0xPPJ5KGvZK74mtZgZ03cdgvr5uj2eoLERfdiSntciHFovFaDQh9s6HbEjtC9MO0/sZo3EShyZK2F9Rel9sNTOg6wiD+fq6qfYhLsLs6+76AMwwKLHuYBOx4ZSK3mIKPvc1/xGHtIYeOogr7I3x9/ctvreaGdB1yGC+maFT7YxnxiDFRuq1tuK3PrGRQ2o6TlMj7Z77cXrclLxdsW9iHcZY1fP7lN4bW81+sNgBUnPt2rXmpZdear9qmkUwbxat+mbRA22/o+iVV15pPvroo+Vj9tEiyOy0r/hbp0+fbj799NP2O01z4cKF5tatW83iAt9+p17vv/9++6hpFo2e5o033mi/apqzZ89mf5zG97doxLWPhsF5/eWXXy4fLxo/zaLhvnysArWBXRXrLgFpz/x4Q6fa6XHGv8fjElLIQ4nzpumNx94s+yrHFeG6YgZiyMJH9lU8tjjWSpTeH1vNDOiVo9gonvCkeA3mxxsy1d6dv066tbTxzV3FoSDGz+Pd5XJeRCaK4+dDFVaC8zn93TEq5+civUe2mhnQK0ZPM15IKL5yvPZ4Q1a1d2sWaByUMi94SLGugAZQbIAOGfz2JWYchhw/j9Pgdj1W5y69T7aaOYZeqa+//rp59tlnm7/85S/Lrx977LHm5s2bjtce4/bt20fqDH7/+98vx3O3cf78+SN/a9HrX46X8znoe4z9prqCRVBqvvnmm+a7775bfs1Y87b7f07GGD/n/L58+XL71UE9Rgn7SsczoFeIIE4wT4UyBHEK4Awmx/vtb397GEy4OF66dGn5eFP8nevXr7dfHRQgWvzWr1sM99Zbb7VfNc3vfve79lHePvjgg/ZR0zz11FPto93QWCSog/PaQrhKtD11VYLxXore+OjZSMXVvGDJuoZItbPvu3PMa7iH+S7i/op3/ONxKeKw1xBp8VhEyFbD+R3fb80M6JXp3rGrhDstjW2IqvZuQ4qtlIKusbDP4n6Pj0tZAnfo8XP22ZOh5oAGUQ3S+2WrmQG9It0irJLWch7TrlXtFBrGiywbldo6XlzuNW4EvlKyGnGWwxDBN/49GkAlTOlbR3rPbDVzDL0SjD3GIiyKZLYdA67JrgvIUK/wzDPPHFkwhvHMxYW3/Uqr/OEPf2gfHfXCCy9kv5BMMuT4ebcQbtFgXy4ko4q0gV0Fi8tkstWShtsVY5Exzbtpqp0paHHJTTYXjFlfd9+lraRe55Dj57HegIxQTdL7ZquZAb1wXCRiUHLhmPWwj2IR1qapdsbduwGp1FW6xsD+i/subSU1RoccPy/95isnide4mq9vptwLxrQ0pqelqVaLAOX67GtinjjzzsH+unHjxtr7jfT6L37xi8Npgfw79jupYq3nj3/8Y/voqJKmXw01/5zz++LFi+1XTXPu3LnlugY1iUMLabpejQzohUpzzeNcVBeOWQ/1BnG+8+uvv75sDK2D8XbGzNOCPSmYL3qWy6+1nr7x8yeffLKoxVGGGj/v3nyFsfPaGNBbbU9dBSHlFCuzSecNMb+1Box7x3HNTVK81CosArj7fUfMCoj7MW2klUsSj7Ntj5PulMpah3XilNCa19Wwh16g559//khlNunikno2YyF1yb6Ly+FS1b6Od955p3nuuecOhzfIhLD6m/t9c6Si035M6IGxol4pGJZJxxnHyrbHCUNDaV+Rtq91WMce+gEDemEYS3vvvffar5plQCrpQjgmUpepIQQaQlxsT0J6noZAurDSEFj0uAzmW+pLt7NuPo3TUgwxfs5xl/4O+4ahoVrF8zQ1lGpkQC/ItWvXllvCWNqZM2far3QcLozMOU8ovlqnsIgGFGuzJ4y1e5OV3XQL4ghWpR3Hu46fE7S6N19Zt86jRD/84Q/bR3UHdMfQC8ESrnycaXMlsvUxr/nUqVOH+46pfSehTiHO+2Vb9Mi9/eyOuuuQs5V4LMfjbZvx83Pnzh3+e6ZH1jxVC9QOpP1R85LK9tALQI8m9hKpqHYlsvWx79K4G6k7Uu3H4XeZlhaHNtjni2C0Vopeq/VNVzt79mz7qAz0INPxRvZh06EZhoXi3fpItZc0HLGNeN516y9qYkDPHMU1sRiL8biTApK+xxBFDCLUHMQCmy729+nTp48s5Uq60/n9w/jTn/7UPjpAQ6m0VHKs09jmvVEIl7B/rJE5GtAtilOWmHvKnOcUzGnpG1jWx37rrm/PBXIVeuT0zGPvit6R2ZBhsF9jsENpvXOkBYuwae+cBmhqTKbjT1a5Jwb0TKVgngpAKMIimJvyXR9TzdL+44Jw3IIcFMx1p6WxUA+rcmkY3XQ7wa7E3uef//zn9lHT/PSnP20fnYxA5c1X+plyP2BAzxBBiOCSVodKwcXK6s288cYb7aOm+d3vfteb2eDiwBh7vJCynxkv32W5Tt2vm24vsXeOOFyzScr91VdfPZKNI6OkA7Fhk66LVWqL42aJlaGoOKbKtZb7+p6kuwrcIghVdyOGIVBZHPdh3/FFxTrHX/o9Nva9lezD47heNEwP9/PiAl1s5XZ8n+te1/g9jtP070pbNW8IHDNp/9R6js46oC9ar4cfEAfzyy+/3P6kTlzg4hKHbDUvc7gLGolpH/66Z3nXvrulMR2m9ulBY7l169aRfX3p0qX2J2UhMKf3SGBfVzxeF73z9ruK4vnKEs41mnXKPRaMkGoi7RlXWKoN1a3dimwrXDfHsRRvvtJN7XKMxbulgfFK9ndfWl67i6vDsY8ZAinRNhXujJ3HaWol3XFuSHHIsdbCuFkHdJZ77GI8s8aiB1Yki0GIAOPtOLcTi+G4CMRGERdO7lKXfk5wYRrgose4/FrjiA1VPo9Si71ihfu6tzjtjp0fNxOjZrEwLp2/tcmmh57Qa4r3/q3Biy++eGRJV4phDDDbi8VwqXfOBZMMSLzZBUGF4jeXzx0X53QMdCX3QD/77LP2UdP85Cc/aR+tZu98fTGg20PPCAd4Lal3euVxrjTBxXnP26PCOKU96X2T5eDkZwpgvHDSmCSY9zUqNSwyJgn7u+R9vukcdHvn63v00UfbR/bQsxBTozWk3pn7HJd05f0zjqvtvf322+2jg/1J75CV3+LYJo0mgnkck9M4OIdfe+219qum+c1vftM+KlMM6CeNods730ysbzGgZ4BgltIqpafeKQCMc59pnTOWa1HWbmIdAsGEnnlMz5H9cD9Ph9X34v4vuQdKMI/DOTFF3Mfe+WZi3cVXX33VPqrL7AN6vLDyOKabab3GG2SUgl55vJUnC5hwS86TLgA6WWy5U4gVL7DsYxfrmBZBKyEjUnJWJC4oc1JBnL3zzcWAns7r2sw+oMdxJlq4jHnGlirBL04vyhkHISvAxV4kaWFWgTOYj4eL6yeffOLKbxNjmCMGudJ7oLEg7qR0u73zzVkUl1lATyc/qffUkqfH9fzzz2ffIuN9EMxjxoHGC8Hc9O9wuhkfpv8xXh5b95pG7J3j7/7u79pHZYrj58et4W7vfDvxHDagz1Sc2pFuakBLLI5z0tKP4825IZgzlhvn4pL6tQBueOlmKqlX7vS//eCCGxuvnMvrzsvO1boFcfbOtxN76LWm3Ge99CtY2pSXybY44dvvHrh69erhz9gWAbD9ST5YojAuccu2aJG3P5XKFJcyZVsErfYnZWLJ4Ph+Vy0h7Jrtu4n7uEbZjaFH9GJj65VFQbq/M2e8Vnrm8TXTK2chGalU9J5inQhqSrfTO49DP5G9893Eosoap67NPqAzLpJSKXxA3QK4OJ7OiRDvWT1n1APE9cI5wRlGcDlXlS4uvZuUXpAYi/9iJyXidxw7300cyoj7vBazD+iIJ0B3hbjueDotYXrqc8ZYOT3zdFHjtb/77rsuMaoqxKV3wfld8nQ1xAr3voI4OiFxsSzqCeydby7GCgP6TD311FPto6b54IMP2kff4+CnWjkhnddN6c0FhUBkEVIwp0FClXVcBU8qFRfZuCofSk+3o5ty72LdiRSAaOBbELud2Fj6+OOP20f1yCKgx3TcqlZXDuPppNPikAC9EhYzWZWCk0oTl95Napj/Hxsx3YDONS1O4SPV3hf0dbJ4LZ3b9X8KD1AZ1z6eLQLgI488chgIv/322yNTFBJ6vazLncalOSmYmrSqAGUqvC4aGPEmFCmYl55qlBLO3x/96EdHxs85jzmfS0bA5rqE7vtln1BLkzoqZBvJ2Gl7xIp0jN29e3fv1/8pZdFD5wOJc1RX3WltjuPptMyfeOKJ++4oRUPDYK6a9BXD1TDUFK9X3WyEqfbhxezGqoxuqbII6PjlL3/ZPuofR0/6xtP3NQ2M56X1HVctYmiAFnhfhkEqWbcYDvG8LlW8XsV6IFPt44idv269RvGWs9EzsOjRHi4YsOjhtt9djYUq0u+zTbnoDItDLA6qI8+/CODLRXKkGsXzN24srFQ6zv30ftkPYGEZrmPp+1wvNAyu9Wm/njt3rv1uHbLpoZOmTr1aWrbd1F0XqfeYzmNKyBR3ZuM5SLHHliEtxs8//9xKdlWrrxiOc7r0Yad4reL6xXuGqfbxpH2M2nro2QR0xJP/pMX3OUkI6vHD5SYuq8bfd0VxC+P1cUoaSKN58w/VjHOjbxppDdXtfePnptrHxTU/1lGlYuoaZBvQ17llKi1i7laW/h0fLAE3tYyHwt+jijWu8kQAJ5C7jKtq11cMhxrmn3fHz7kGdReQ8QZBw0sNJPZzTdPXsgro29wej3/D9LD0b7mwPPvss4PdQ/3atWvLwrd40DAfnhR7LM6QatVXDEcPqobzo9tDN9U+jXhs1ZR2zyqgP/roo+2j9XroCT10euppDJ7GQFx6dRupYXDx4sXD1jYn6Ouvv75cxjU9l1QzglffBZXgltKipeK9p2tMuh6Yap9GXDEuLrtbumx76F999VX7aD2Mq8Q56jQItg3qtLopfIv3L+fEZG55ut+2pP5iONQwXS32zukxmmqfTqydGnqIdc6KLYrrQ5U5QT3hg6ZQLp1k67h8+fKyIRCfnyBOMLe1LR216p4KNcz4iOPnXGtSYDHVPj4DegY2LYrrw/h2PJnoZdNyPgnPx1g5Y2AJaTTS66TZS08fSpuih9qXAeM8rqHxG3vosQNgqn18XI9TUKfDVktQL74org/3HOekSqjCZSx8FX5OFXvf3HJvcSj1iwEtqmG6GgFkVWPGVPs0auylZxXQaXXFavVdKtWZThbHu6lW7y48Q8uO3jtp+XhyOrdcOtmf/vSn9tFRNUxXW9WYcXGp6dRYGJdVQEecjrDrPc9JlZ85c6b96mA1udRIoEVHrzw+BwHcueXSyWgAr5ouVEMPfdX9Jv7+7/++faSxxR56LVPXsrh9akQvmsVhQPrqiy++WD7eFr1wKtZTIOcgYCPNHovl0ti709Gkk3H+kNnqIpizLkTp4i08E64dd+7csd5mIux/Pgew70u/TS+y66GTskqpboLwqtTWulgQJvb66ZnTK0/BnJPPueXSZlal2+Pdxkq1avyca5fBfDpcr1PxIZ9HDSvGZRfQOSFimnzVPNeTUK3+ox/9aJlWpzfRh4PBueXS5moeQ1713muYez83cTZBDYVx2QV0/OY3v2kfHaT2Ymr8JJxspNiZT35cpTwNB1KDTi+RNsOFs69glcxazIaVqm/8nOuJM2Km97Of/ax91DQff/xx+6hcWQb0NM4NgvmqHnZE8GZMj0VhYuolnWik1WPlOn+Xu6dJ2syqHirTRWvQV4BFQ8Yhu+nFwjh76DMWe+knpd2ZkkavPAZ+Tq6XX365uXv37nJ8nLQ6Jx2BPaEAj38raX1/+MMf2kdH1TBdjY5DX+avhvc+R7UF9IYq9xzduXPn3qJ3TYX+cvv222/bnxx15cqVw99J25kzZ5b/fpULFy4c/i7Psei5tz+RdJxFA/nIeZm2xx9/vP2Nst28efO+9872xRdftL+hqZ06daqazyHbHjqp8dj66ktzMU7+0ksvtV8dFEgwLs567im13oeee0zpx5sqSFqNdHvfuVLL/Ou+XiDXnbhstabVncVUsmwDOlZ9UIyRd9ddp7qWivV1FrVgXD3emY2/53i6dLJV09VqGT/vW5HMxWT2K64YV3phXNYBvftBMdeQVdy6664TzBknTwF6HbSq43g6c9NdIU46XrylcELDu5Yeal8PsIapenNW0zh6divFRXw4BG+kCtK4oAMBnHXXd7kZAqvSxTXe+XsGdul+TFX78Y9/3H71PRrGNazlwFDDQw891H51YIjVLLWbeFwSJ0peMS7rHjotr9TrJpDHYE6vgBT7rnc2IvUeW9iMyRvQpfv1TVfj/Kxl/nVf7y/OxtF+0KiKHb5d7tQ5d1kHdHRb/nx4rLnOnPIhFoXhgkS63qAuHa9vuho1K8cVoJakL6C7yuQ8xHqrkm/Ukn1ApyI9LQVLT4B7lA9dgGNQl45HurnvQllTD7VbEMf1opbGzNzFzl3J4+jZB3SCLWlx7mK0aeHbJgzq0mrXr1+/L5XJOVPTcqfdQGG6fT7iErAl3xs966K4faAnQqFcrOa1UE4145yg6Kgb0Mmc0diuBQVxaQ4+PXOK4cbqYGgzTD1mtVDw2dABLFH2PfSp2VOXjmIWSF+hUU09VAJGCuagMWMwnw9S7unzsChOR6wK6hcvXjxyUisPZFu4aQ9pY23u1VdfbR99j17QOos4lcJ0+/zVcG02oG+pL6hzIxdSj32La2i+WNqXKVc0yLQZeud9RUaMndfUQ43jskynjYuZSFMxoO8gBfVY+EM659lnn11uJad2SkEwSp9TTQFoKK+99lr76KjaeqixEX/27Nn2kTQtA/qOUlBn7nucosIJTm89riev+YmLodSUIh4CjaG+xWRYCyLO+y0d+yFmKTyOtC8G9IEw95058HEhCcZsuOMb1ZUlL2aQsw8++KB91DRPPfVU+0jroG6kT22987fffrt9dNCYGWJBK2kbBvQBsbwg61azSl0cQ0t3f2OsNi5Pq/2zh74djul4j4OoljurJe+88077yHR7ybhWMJQaP++5MaCPIK0jzyp2aQ1hcMc20vD8X/vHCZoaWHxOFjKt74033mgfHUWjiF5qLbo1GLU1ZmpBA5ZgzlAqHbO5MqCPiBvDkIaPRXMEEA4IeuwcJNqf2DuvaUWzXRHAVk3xqzndTkPepV7nKxa9bnLtZej0+eefP5z2Nuf6EAP6yDjBKZq7efPmkZ4LY+qMrTNVyjT8fsSbiTh+vj7mnffN6eWCme6rUIuYbTPDM29xivGqDFMf6qBS0SPHOAXQc2VAnwgHE711eu2xpZjmrvP/GhY+mAt6mbEy2R76+uLYeeyR1rg6WmyMx/XCNT8xe8Q4+DrXW7J4XJsTaqTmPKRkQJ8QFzvG1RlfjwVYXBToqdNjn3PBRUnivGFSaLHWQatxEfzyyy/br44GtBpXR4sNGHvo80ajPTVAadCvKupM+B1S7QkN1rnXSBjQ94BpLbdu3Vqm4mNrjwslBxDj63F8V8OL6fZf/vKX7SOdpDv2mHo5BLMaZwn8/ve/X/6fDJzT1eYvzkI4Ke1+/vz5w4JHrtP0zmePu61pv65evXpv0XLkrndHtkWL8t7nn3/e/paGcuXKlSP7+cMPP2x/opMsGqFH9p37UDm5c+fOkeN20bFqf3JU9xqx6vfmxoA+E99+++29S5cu3XvwwQePHEhsZ86cubdoHRrcB9A9URc9q/YnWsfLL798ZP+xvfDCC+1PpfnjepqO3SeffPLe3bt3258c6F4j+DoXBvSZoQXJBTIeUHGjJ2+A3w6ZkLgvCebdk1nHO3fu3JF9+PDDDy+PWSkXHK+x48R1Ielr8Od0jTCgz9Qnn3xy7+mnnz5ycPVtBvj1kBKOJ7HBfDvdYzJeDKVcxMY91wWCfO7BHA/wn8WL10wxtYoCOdYcjyubrUIVZyrOoVDphz/84fJxdyU0HtdS2c0+O3369GF1NlXtFCXGCmWt55FHHjk8BjmGmLEh5YZiToqP49TVaBHMl0XLuV0jDOiZ2TTAryMFew5e5tJS0clWStB/7rnnDqeo8B4JQlYkb+ehhx46rGznguf8feWK6+czzzzTfvW9XIM5DOiZGyPARxzU9GgJ8I8++ujh8paxtz9X7AvuCBYXhmCVJ9fb3h6rZnFLYC56rH4o5ezFF188svIhc825RuQYzGFALwxp5bjwB0E++eqrr1b+bBsc9OlEyAGBfM7LNkrSLgzoWkoNAdaY/+abb5Y9fxYRSQsr5M5xc0mlM6DrWPTACfIp4H/88cf3rYM+ZwwVXLlyxTS7pOIZ0CVJKoBruUuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQUwoEuSVAADuiRJBTCgS5JUAAO6JEkFMKBLklQAA7okSQV44N5C+1jJ//rvTfNfX2yav/yf9hsbePhfN81f/9um+Tc/bZpTjx88ZpMkaUQG9D7X/v12wfw4j/1tG9x/vAj0TxwE/H/2z9sfSpK0GwN6n//xX5rmv/3n9osZIgvABhoKSN+jkUBjQZJUFQP60L7+vGn+8X8vti+a5sv/efD4//1j+8OJkRH4q79umgf/5UH6H/yfrx/8FweZAklSEQzoU/ju/y4C/e2m+YdPDwI9Qf4fPmt/OAPd9P+qXn7MDEQOH+SPxucqDD8dNwS1y79VHjjvn/6PTfPkc+03NEcG9BzFi2S6mPZ9T5KGQqP9P82oI6L7GNBLZuCXNJR/d6Zp/sOL7ReaIwO6+sf50zBBV18jgN/j95Wv44ZN/upfHczOWOXYf/vXTtuUJmJAlySpAK4UJ0lSAQzokiQVwIAuSVIBDOiSJBXAgC5JUgEM6JIkFcCALklSAQzokiQVwIAuSVIBDOiSJBXAgC5JUgEM6JIkFcCALklSAQzokiQVwIAuSVIBDOiSJBXAgC5JUgEM6JIkFcCALklSAQzokiQVwIAuSVIBDOiSJBXAgC5JUgEM6JIkFcCALklSAQzokiQVwIAuSVIBDOiSJBXAgC5JUgEM6JIkFcCALklSAQzokiQVwIAuSVIBDOiSJBXAgC5JUgEM6JIkFcCALklSAQzokiRlr2n+P5H8M38lCNoIAAAAAElFTkSuQmCC"

        imgdata = base64.b64decode(photob64.split(',')[1])
        bz = io.BytesIO(imgdata)
        doc = docx.Document(template_file_path)
        for i, p in enumerate(doc.paragraphs):
            # print(str(i) + ":" + str(p.text))
            if len(p.text) != 0:
                for i in range(len(p.runs)):
                    pass
                    # print(str(i) + ':::::')
                    # print(p.runs[i].text)
            if 'd.inspector_signature_image' in p.text:
                current_app.logger.info(f'FOUND IT!')
                p.runs[-1].add_break()
                # p.runs[-1].add_text('******************')
                p.runs[-1].add_picture(bz)
        doc.save(template_file_path)             # Save file

        # files = {'template': (file_name, file.read(), mimetypes.guess_type(file_name))}
        # resp = requests.post(url=cls.document_generator_url, files=files)
        # if resp.status_code != 200:
        #     current_app.logger.warn(f'Docgen-api/push-template replied with {str(resp.text)}')
        #     return False
        # return True